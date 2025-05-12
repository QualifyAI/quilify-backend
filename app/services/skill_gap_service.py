from typing import Dict, List, Optional, Any
from fastapi import UploadFile, HTTPException, status
import io
import PyPDF2
import docx

from app.db.repositories.skill_gap_repository import SkillGapRepository
from app.services.skill_gap_ai_service import SkillGapAIService
from app.models.skill_gap import SkillGapAnalysis
from app.services.resume_service import ResumeService
from app.utils.resume_parser import parse_resume_file

class SkillGapService:
    def __init__(self):
        self.repository = SkillGapRepository()
        self.ai_service = SkillGapAIService()
        self.resume_service = ResumeService()
    
    async def analyze_skills(
        self,
        user_id: str,
        resume_text: Optional[str] = None,
        resume_file: Optional[UploadFile] = None,
        resume_id: Optional[str] = None,
        job_description: Optional[str] = None,
        job_posting_url: Optional[str] = None,
        github_url: Optional[str] = None
    ) -> SkillGapAnalysis:
        """
        Analyze skills gap between resume and job description
        """
        # Get resume text (multiple sources with priority: resume_id > resume_file > resume_text)
        if resume_id:
            # Use a stored resume
            resume = await self.resume_service.get_resume(resume_id)
            
            # Check if the resume belongs to the user
            if resume.userId != user_id:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Not authorized to use this resume"
                )
                
            resume_text = resume.content
        elif not resume_text and resume_file:
            # Parse the uploaded file
            resume_text = await parse_resume_file(resume_file)
        
        # Validate resume text is available
        if not resume_text:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Resume text, file, or ID must be provided"
            )
        
        # Get job description (either from provided text or by fetching from URL)
        processed_job_description = job_description
        if not processed_job_description and job_posting_url:
            processed_job_description = await self.ai_service.fetch_job_description(job_posting_url)
        
        if not processed_job_description:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Job description or job posting URL must be provided"
            )
        
        # Analyze using AI service
        analysis_result = await self.ai_service.analyze_skill_gap(
            resume_text=resume_text,
            job_description=processed_job_description,
            job_posting_url=job_posting_url
        )
        
        # Prepare data for saving to database
        analysis_data = {
            "job_title": analysis_result.job_title,
            "job_description": processed_job_description,
            "resume_text": resume_text,
            "match_percentage": analysis_result.match_percentage,
            "matched_skills": [skill.model_dump() for skill in analysis_result.matched_skills],
            "missing_skills": [skill.model_dump() for skill in analysis_result.missing_skills],
            "project_recommendations": [project.model_dump() for project in analysis_result.project_recommendations],
            "improvement_suggestions": analysis_result.improvement_suggestions,
            "overall_assessment": analysis_result.overall_assessment,
            "job_posting_url": job_posting_url
        }
        
        # Save to database
        saved_analysis = await self.repository.create_analysis(user_id, analysis_data)
        
        # If resume_file was provided and parsed successfully, 
        # consider offering to save it for future use
        if resume_file and not resume_id and not resume_text:
            # This would be implemented in frontend flow
            pass
            
        return saved_analysis
    
    async def get_analysis_by_id(self, analysis_id: str) -> Optional[SkillGapAnalysis]:
        """
        Get a skill gap analysis by ID
        """
        return await self.repository.get_analysis_by_id(analysis_id)
    
    async def get_user_analyses(self, user_id: str) -> List[SkillGapAnalysis]:
        """
        Get all skill gap analyses for a user
        """
        return await self.repository.get_analyses_by_user_id(user_id)
    
    async def delete_analysis(self, analysis_id: str) -> bool:
        """
        Delete a skill gap analysis
        """
        return await self.repository.delete_analysis(analysis_id)
    
    async def _parse_resume_file(self, file: UploadFile) -> str:
        """
        Parse resume file (PDF, DOCX) to extract text
        """
        content = await file.read()
        file_extension = file.filename.split('.')[-1].lower() if file.filename else ''
        
        try:
            # Parse PDF file
            if file_extension == 'pdf':
                return self._extract_text_from_pdf(content)
            
            # Parse DOCX file
            elif file_extension == 'docx':
                return self._extract_text_from_docx(content)
            
            # Handle plain text files
            elif file_extension in ['txt', 'text']:
                return content.decode('utf-8')
            
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, or TXT file."
                )
                
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Error parsing resume file: {str(e)}"
            )
        finally:
            # Reset file pointer for potential future reads
            await file.seek(0)
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF content
        """
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        # Extract text from each page
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
            
        return text
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX content
        """
        doc = docx.Document(io.BytesIO(content))
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
