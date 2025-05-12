import io
import PyPDF2
import docx
from fastapi import UploadFile, HTTPException, status


async def parse_resume_file(file: UploadFile) -> str:
    """
    Parse a resume file (PDF, DOCX, TXT) and extract its text content
    """
    content = await file.read()
    file_extension = file.filename.split('.')[-1].lower() if file.filename else ''
    
    try:
        # Parse PDF file
        if file_extension == 'pdf':
            return extract_text_from_pdf(content)
        
        # Parse DOCX file
        elif file_extension == 'docx':
            return extract_text_from_docx(content)
        
        # Handle plain text files
        elif file_extension in ['txt', 'text']:
            return content.decode('utf-8')
        
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, or TXT file."
            )
            
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Error parsing resume file: {str(e)}"
        )
    finally:
        # Reset file pointer for potential future reads
        await file.seek(0)


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF content
    """
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
    text = ""
    
    # Extract text from each page
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
        
    return text


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX content
    """
    doc = docx.Document(io.BytesIO(content))
    text = ""
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        text += para.text + "\n"
        
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
            
    return text
